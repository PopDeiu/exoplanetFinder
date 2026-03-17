#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pentru generarea documentației în format .docx (Word)
Utilizează python-docx pentru a crea document profesional
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Crează documentul
doc = Document()

# Setează style-uri
style = doc.styles['Normal']
style.font.name = 'Segoe UI'
style.font.size = Pt(11)

# --- TITLU PRINCIPAL ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('📚 DOCUMENTAȚIE PROIECT')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(44, 62, 80)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('EXOPLANETFINDER - Aplicație de Explorare Datelor și Analiză')
run.font.size = Pt(14)
run.font.italic = True
run.font.color.rgb = RGBColor(127, 140, 141)

doc.add_paragraph()  # Spațiu

# --- CUPRINS ---
doc.add_heading('Cuprins', level=1)
toc_items = [
    '1. Prezentare Generală',
    '2. Tehnologii Utilizate',
    '3. Structura Proiectului',
    '4. Componente Principale',
    '5. Pagini și Funcționalități',
    '6. Instalare și Configurare',
    '7. Ghid de Deployment',
    '8. Procesul de Configurare',
    '9. Arhitectura Aplicației',
    '10. Note de Dezvoltare',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# --- SECȚIUNEA 1 ---
doc.add_heading('1. PREZENTARE GENERALĂ', level=1)

doc.add_heading('1.1 Despre Proiect', level=2)
doc.add_paragraph(
    'ExoplanetFinder este o aplicație web interactivă construită cu Streamlit, '
    'concepută pentru explorarea și analiza datelor din misiunile spațiale TESS și Kepler. '
    'Aplicația oferă un interfață user-friendly pentru:'
)

items = [
    'Căutarea și analiza stelelor după identificatori (TIC, TOI, KIC)',
    'Explorarea sistemelor cunoscute cu candidați și planete confirmate',
    'Analiza și identificarea semnalelor false pozitive',
    'Descoperirea țintelor potențial netestare',
    'Ajustarea parametrilor de analiză în timp real',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.2 Obiectivele Aplicației', level=2)
objectives = [
    'Furnizare de acces ușor la datele publice din TESS și Kepler',
    'Crearea unei platforme interactivă pentru analiza curbelor de lumină',
    'Facilitarea procesului de descoperire de potențiale exoplanete',
    'Generarea rapoartelor în format PDF pentru rezultatele analizei',
    'Oferirea unei interfețe configurabile pentru parametrii analizei',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('1.3 Público-țintă', level=2)
audience = [
    'Pasionații de astronomie',
    'Cercetători care lucrează cu datele TESS/Kepler',
    'Studenți și doritori de aprofundare în analiza datelor spațiale',
    'Oricine dorește să exploreze datele exoplanetelor',
]
for aud in audience:
    doc.add_paragraph(aud, style='List Bullet')

# --- SECȚIUNEA 2 ---
doc.add_page_break()
doc.add_heading('2. TEHNOLOGII UTILIZATE', level=1)

doc.add_heading('2.1 Technology Stack', level=2)

# Tabel cu tehnologii
table = doc.add_table(rows=9, cols=4)
table.style = 'Light Grid Accent 1'

# Header
header_cells = table.rows[0].cells
header_cells[0].text = 'Componentă'
header_cells[1].text = 'Tehnologie'
header_cells[2].text = 'Versiune'
header_cells[3].text = 'Utilizare'

tech_data = [
    ['Frontend', 'Streamlit', 'Latest', 'Interfața web interactivă'],
    ['Runtime', 'Python', '3.9+', 'Limbajul de programare'],
    ['Data Science', 'Pandas, NumPy', 'Latest', 'Manipulare și procesare date'],
    ['Astronomie', 'Lightkurve, Astropy', 'Latest', 'Procesare curbelor de lumină'],
    ['Queries', 'Astroquery', 'Latest', 'Interogări baze date astronomice'],
    ['Grafice', 'Matplotlib, Plotly', 'Latest', 'Vizualizare date'],
    ['Rapoarte', 'ReportLab', 'Latest', 'Generare PDF-uri'],
    ['Containerizare', 'Docker, Docker Compose', 'Latest', 'Deployment și orchestrare'],
]

for idx, row_data in enumerate(tech_data, start=1):
    row_cells = table.rows[idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row_cells[col_idx].text = cell_data

doc.add_heading('2.2 Dependențe Python', level=2)
doc.add_paragraph('Dependențele principale sunt definite în requirements.txt:')

dependencies = [
    'streamlit - Framework pentru interfața web',
    'lightkurve - Procesare și analiză curbelor de lumină',
    'matplotlib - Grafice statice 2D',
    'astropy - Unități și timeseries analysis',
    'pandas - DataFrames și manipulare date',
    'numpy - Operații numerice',
    'astroquery - Acces la cataloagele astronomice',
    'plotly - Grafice interactive',
    'streamlit-plotly-events - Integrare Plotly cu Streamlit',
    'reportlab - Generare documente PDF',
    'pillow - Procesare imagini',
]
for dep in dependencies:
    doc.add_paragraph(dep, style='List Bullet')

# --- SECȚIUNEA 3 ---
doc.add_page_break()
doc.add_heading('3. STRUCTURA PROIECTULUI', level=1)

doc.add_heading('3.1 Organizarea Directoarelor', level=2)
p = doc.add_paragraph()
p.add_run('exoplanetFinder/\n').font.name = 'Courier New'
p.add_run('├── exo_app.py\n').font.name = 'Courier New'
p.add_run('├── requirements.txt\n').font.name = 'Courier New'
p.add_run('├── Dockerfile\n').font.name = 'Courier New'
p.add_run('├── docker-compose.yml\n').font.name = 'Courier New'
p.add_run('├── pages/\n').font.name = 'Courier New'
p.add_run('│   ├── 1_Search_for_a_Star.py\n').font.name = 'Courier New'
p.add_run('│   ├── 2_Explore_Planet_Candidates.py\n').font.name = 'Courier New'
p.add_run('│   ├── 3_Explore_False_Positives.py\n').font.name = 'Courier New'
p.add_run('│   ├── 4_Find_Untested_Targets.py\n').font.name = 'Courier New'
p.add_run('│   ├── 5_Settings.py\n').font.name = 'Courier New'
p.add_run('│   ├── 6_Star_explorer.py\n').font.name = 'Courier New'
p.add_run('│   └── 7_TESS_Planet_search.py\n').font.name = 'Courier New'
p.add_run('├── utils/\n').font.name = 'Courier New'
p.add_run('│   ├── __init__.py\n').font.name = 'Courier New'
p.add_run('│   ├── ui_styles.py\n').font.name = 'Courier New'
p.add_run('│   ├── settings_manager.py\n').font.name = 'Courier New'
p.add_run('│   ├── data_fetchers.py\n').font.name = 'Courier New'
p.add_run('│   └── analysis_engine.py\n').font.name = 'Courier New'
p.add_run('└── assets/\n').font.name = 'Courier New'

doc.add_heading('3.2 Rolul Fiecărui Fișier Cheie', level=2)

doc.add_heading('exo_app.py - Pagina Principală', level=3)
doc.add_paragraph('Fișierul de intrare al aplicației. Inițializează:')
doc.add_paragraph('Configurarea paginii Streamlit (layout, title)', style='List Bullet')
doc.add_paragraph('Stilurile și temele UI (fundal cosmic)', style='List Bullet')
doc.add_paragraph('Session state pentru gestionarea datelor în sesiune', style='List Bullet')
doc.add_paragraph('Încărcarea setărilor persistente din fișier', style='List Bullet')

# --- SECȚIUNEA 4 ---
doc.add_page_break()
doc.add_heading('4. COMPONENTE PRINCIPALE', level=1)

doc.add_heading('4.1 ui_styles.py - Stilizare și Teme', level=2)
doc.add_paragraph('Module responsabil pentru styling-ul UI al aplicației:')
doc.add_paragraph('set_galaxy_background() - Aplică fundal cu gradient cosmic', style='List Bullet')
doc.add_paragraph('set_sidebar_style() - Stilizează sidebar-ul', style='List Bullet')
doc.add_paragraph('apply_pro_plotting_style() - Configurează tema Dark Mode pentru Matplotlib', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Gradienți disponibili: ').font.bold = True
p.add_run('default, nebula, void, cosmic, stellar')

doc.add_heading('4.2 settings_manager.py - Gestionare Setări', level=2)
doc.add_paragraph('Gestionează salvarea și încărcarea setărilor persistente ale aplicației:')
doc.add_paragraph('load_settings() - Încarcă setări din fișierul JSON', style='List Bullet')
doc.add_paragraph('save_settings() - Salvează setări în fișierul JSON', style='List Bullet')
doc.add_paragraph('get_setting() - Obține o anumită setare', style='List Bullet')
doc.add_paragraph('update_setting() - Actualizează o setare', style='List Bullet')

doc.add_heading('Setări Persistente Disponibile', level=3)
table2 = doc.add_table(rows=6, cols=4)
table2.style = 'Light Grid Accent 1'

header_cells = table2.rows[0].cells
header_cells[0].text = 'Setare'
header_cells[1].text = 'Tip'
header_cells[2].text = 'Valoare Default'
header_cells[3].text = 'Descriere'

settings_data = [
    ['bin_size', 'int', '10', 'Dimensiunea binning pentru agregare date'],
    ['sigma_val', 'float', '5.0', 'Prag sigma pentru outliers'],
    ['period_range', 'list[float]', '[1.0, 30.0]', 'Interval perioade de căutare (zile)'],
    ['selected_missions', 'list[str]', '["TESS", "Kepler", "K2"]', 'Misiuni selectate'],
    ['selected_authors', 'list[str]', '["SPOC", "Kepler"]', 'Sursele de date preferate'],
]

for idx, row_data in enumerate(settings_data, start=1):
    row_cells = table2.rows[idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row_cells[col_idx].text = cell_data

doc.add_heading('4.3 data_fetchers.py - Preluare Date', level=2)
doc.add_paragraph('Gestionează interogările la baze de date astronomice externe:')
doc.add_paragraph('get_toi_catalog() - Descarcă catalogul TESS Objects of Interest (TOI)', style='List Bullet')
doc.add_paragraph('search_toi_catalog() - Filtrează catalogul TOI cu criterii', style='List Bullet')
doc.add_paragraph('fetch_star_data() - Preia date despre o stea', style='List Bullet')
doc.add_paragraph('get_star_parameters() - Obține parametri stelari specifici', style='List Bullet')
doc.add_paragraph('fetch_catalog_targets() - Descarcă ținte din cataloage', style='List Bullet')
doc.add_paragraph('fetch_untested_targets() - Identifică ținte potențial netestare', style='List Bullet')

doc.add_heading('4.4 analysis_engine.py - Motor de Analiză', level=2)
doc.add_paragraph('Conține logica principală pentru procesarea și analiza curbelor de lumină:')
doc.add_paragraph('process_selected_data() - Procesează datele selectate', style='List Bullet')
doc.add_paragraph('generate_pdf_report() - Generează raport PDF cu rezultate', style='List Bullet')
doc.add_paragraph('Utilizează Box Least Squares (BLS) pentru identificare perioade', style='List Bullet')
doc.add_paragraph('Efectuează normalizare și filtrare date', style='List Bullet')

# --- SECȚIUNEA 5 ---
doc.add_page_break()
doc.add_heading('5. PAGINI ȘI FUNCȚIONALITĂȚI', level=1)

pages_info = [
    ('1_Search_for_a_Star.py', 'Căutare și Analiză Stele', [
        'Căutare stele după TIC ID, TOI ID, sau nume',
        'Descărcarea curbei de lumină din arhive',
        'Afișare date stelare (magnitudine, temperatură, rază, masă)',
        'Grafice interactive cu curba de lumină',
        'Export rezultate',
    ]),
    ('2_Explore_Planet_Candidates.py', 'Explorare Candidați', [
        'Explorare sisteme cu candidații de planete confirmate',
        'Filtrare după parametri (rază planetă, perioada orbitală)',
        'Afișare detali despre fiecare candidat',
        'Comparare multiple sisteme',
    ]),
    ('3_Explore_False_Positives.py', 'Analiza False Positive-uri', [
        'Explorare cazuri de false positive în date',
        'Analiza semnalelor care simulează planete',
        'Învățare despre erori comune în detecție',
    ]),
    ('4_Find_Untested_Targets.py', 'Descoperire Ținte Netestare', [
        'Identificare ținte potențial netestare',
        'Filtrare după criterii de oportunitate',
        'Prioritizare pentru noi investigații',
    ]),
    ('5_Settings.py', 'Configurare Aplicație', [
        'Dimensiune binning pentru agregarea datelor',
        'Prag sigma pentru detecția outliers',
        'Interval de perioade pentru căutare',
        'Selectare misiuni (TESS, Kepler, K2)',
        'Selectare surse de date',
        'Tema vizuală (fundal cosmic)',
    ]),
    ('6_Star_explorer.py', 'Explorator Avansat', [
        'Navigare interactivă prin cataloage stele',
        'Analiză multi-criteriu',
        'Vizualizare în timp real',
    ]),
    ('7_TESS_Planet_search.py', 'Căutare Specifică TESS', [
        'Căutare optimizată pentru datele TESS',
        'Acces direct la catalogul TOI',
        'Vizualizare specifice TESS',
    ]),
]

for file_name, section_title, features in pages_info:
    doc.add_heading(section_title, level=2)
    doc.add_paragraph(f'Fișier: {file_name}')
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

# --- SECȚIUNEA 6 ---
doc.add_page_break()
doc.add_heading('6. INSTALARE ȘI CONFIGURARE', level=1)

doc.add_heading('6.1 Cerințe Preliminare', level=2)
reqs = [
    'Python 3.9+ - Limbajul de programare',
    'pip - Package manager Python',
    'Git (opțional) - Pentru clonare repository',
    'Conexiune Internet - Necesară pentru accesul la baze date astronomice',
]
for req in reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_heading('6.2 Instalare Locală (Fără Docker)', level=2)

doc.add_heading('Pasul 1: Clonare/Download Proiect', level=3)
p = doc.add_paragraph()
p.add_run('git clone <repository-url>\ncd exoplanetFinder').font.name = 'Courier New'

doc.add_heading('Pasul 2: Creare Mediu Virtual', level=3)
doc.add_paragraph('Pe Windows:')
p = doc.add_paragraph()
p.add_run('python -m venv .venv\n.venv\\Scripts\\activate').font.name = 'Courier New'

doc.add_paragraph('Pe Linux/Mac:')
p = doc.add_paragraph()
p.add_run('python3 -m venv .venv\nsource .venv/bin/activate').font.name = 'Courier New'

doc.add_heading('Pasul 3: Instalare Dependențe', level=3)
p = doc.add_paragraph()
p.add_run('pip install -r requirements.txt').font.name = 'Courier New'

doc.add_heading('Pasul 4: Pornire Aplicație', level=3)
p = doc.add_paragraph()
p.add_run('streamlit run exo_app.py').font.name = 'Courier New'

doc.add_paragraph('Aplicația va fi disponibilă la http://localhost:8501')

# --- SECȚIUNEA 7 ---
doc.add_page_break()
doc.add_heading('7. GHID DE DEPLOYMENT', level=1)

doc.add_heading('7.1 Deployment cu Docker Compose', level=2)
doc.add_paragraph(
    'Aceasta este modalitatea recomandată pentru deployment în producție.'
)

doc.add_heading('Cerințe Preliminare', level=3)
docker_reqs = [
    'Docker Desktop (Windows/Mac) sau Docker Engine (Linux)',
    'Docker Compose (integrat în Docker Desktop 1.27+)',
]
for req in docker_reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_heading('Pași de Deployment', level=3)

doc.add_heading('1. Verificare Instalare Docker', level=4)
p = doc.add_paragraph()
p.add_run('docker --version\ndocker compose version').font.name = 'Courier New'

doc.add_heading('2. Build și Pornire Container', level=4)
p = doc.add_paragraph()
p.add_run('cd c:\\Exoplanete\\exoplanetFinder\ndocker compose up --build').font.name = 'Courier New'

doc.add_heading('3. Acces la Aplicație', level=4)
doc.add_paragraph('Odată cu build-ul completat, aplicația va fi disponibilă la:')
p = doc.add_paragraph()
p.add_run('http://localhost:8501').font.name = 'Courier New'

doc.add_heading('4. Oprire Container', level=4)
p = doc.add_paragraph()
p.add_run('docker compose down').font.name = 'Courier New'

doc.add_heading('7.2 Configurare Docker Compose', level=2)
doc.add_paragraph('Fișierul docker-compose.yml conține:')
docker_config = [
    'version: 3.8 - Versiunea Docker Compose API',
    'services.app - Serviciul principal al aplicației',
    'image: exoplanetfinder:latest - Imaginea Docker',
    'build: . - Build dintr-un Dockerfile în directorul curent',
    'restart: unless-stopped - Restart automat dacă cade',
    'ports: "8501:8501" - Port mapping (host:container)',
    'environment - Variabile de mediu pentru Streamlit',
]
for item in docker_config:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.3 Configurare Dockerfile', level=2)
doc.add_paragraph('Dockerfile definește pas cu pas cum se construiește imaginea containerului:')

# --- SECȚIUNEA 8 ---
doc.add_page_break()
doc.add_heading('8. PROCESUL DE CONFIGURARE', level=1)

doc.add_heading('8.1 Setări Persistente', level=2)
doc.add_paragraph(
    'Setările aplicației sunt salvate în fișierul settings.json în directorul root al proiectului.'
)

doc.add_heading('Structura settings.json', level=3)
p = doc.add_paragraph()
p.add_run('{\n    "bin_size": 10,\n    "sigma_val": 5.0,\n    "period_range": [1.0, 30.0],\n    '
          '"selected_missions": ["TESS", "Kepler", "K2"],\n    "selected_authors": ["SPOC", "Kepler"]\n}').font.name = 'Courier New'

doc.add_heading('8.2 Session State', level=2)
doc.add_paragraph('Streamlit utilizează session state pentru a gestiona datele în cadrul unui sesiune:')
doc.add_paragraph('search_result - Rezultatele căutării unei stele', style='List Bullet')
doc.add_paragraph('explore_planets_results - Datele din pagina de explorare planete', style='List Bullet')
doc.add_paragraph('explore_fps_results - Datele din pagina false positive', style='List Bullet')
doc.add_paragraph('untested_results - Țintele netestare identificate', style='List Bullet')

doc.add_heading('8.3 Cache-ul Aplicației', level=2)
doc.add_paragraph('Pentru performanță optimă, aplicația utilizează cache pentru datele descărcate:')
doc.add_paragraph('TOI Catalog: Cache 24 ore', style='List Bullet')
doc.add_paragraph('Date Stele: Cache 7 zile', style='List Bullet')
doc.add_paragraph('Curve de Lumină: Descărcate din Lightkurve/MAST', style='List Bullet')

doc.add_heading('8.4 Personalizare Teme', level=2)
doc.add_paragraph('Aplicația suportă multiple teme de fundal cosmic:')
themes = [
    'default - Gradient albastru-violet standard',
    'nebula - Culori de nebuloasă purpurie',
    'void - Void cosmic radiant',
    'cosmic - Gradient cosmic abstract',
    'stellar - Gradient conic (stellar)',
]
for theme in themes:
    doc.add_paragraph(theme, style='List Bullet')

# --- SECȚIUNEA 9 ---
doc.add_page_break()
doc.add_heading('9. ARHITECTURA APLICAȚIEI', level=1)

doc.add_heading('9.1 Fluxul de Date', level=2)
doc.add_paragraph(
    'Aplicația urmează o arhitectură standard web cu multiple nivele de procesare.'
)

doc.add_heading('9.2 Componentele Functionali', level=2)

doc.add_heading('Frontend Layer', level=3)
doc.add_paragraph('Streamlit UI components (buttons, sliders, text input)', style='List Bullet')
doc.add_paragraph('Plotly interactive charts', style='List Bullet')
doc.add_paragraph('Custom CSS styling (galaxy background)', style='List Bullet')

doc.add_heading('Business Logic Layer', level=3)
doc.add_paragraph('Data fetching și preprocessing', style='List Bullet')
doc.add_paragraph('Analysis engine (BLS algorithm)', style='List Bullet')
doc.add_paragraph('Settings management', style='List Bullet')
doc.add_paragraph('PDF report generation', style='List Bullet')

doc.add_heading('Data Layer', level=3)
doc.add_paragraph('External API calls (Astroquery)', style='List Bullet')
doc.add_paragraph('Local cache (Streamlit @cache_data)', style='List Bullet')
doc.add_paragraph('Settings persistence (JSON file)', style='List Bullet')

doc.add_heading('9.3 Integrații Externe', level=2)
table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Light Grid Accent 1'

header_cells = table3.rows[0].cells
header_cells[0].text = 'Serviciu'
header_cells[1].text = 'Provider'
header_cells[2].text = 'Utilizare'

integrations = [
    ['MAST Archive', 'STScI (NASA)', 'Curve de lumină TESS/Kepler'],
    ['TOI Catalog', 'ExoFOP (Caltech)', 'TESS Objects of Interest'],
    ['NASA Exoplanet Archive', 'NASA/Caltech', 'Datele planetelor confirmate'],
    ['Simbad', 'CDS (Strasbourg)', 'Identificare stele și ID-uri'],
    ['SkyView', 'NASA HEASARC', 'Imagini astronomice'],
]

for idx, row_data in enumerate(integrations, start=1):
    row_cells = table3.rows[idx].cells
    for col_idx, cell_data in enumerate(row_data):
        row_cells[col_idx].text = cell_data

# --- SECȚIUNEA 10 ---
doc.add_page_break()
doc.add_heading('10. NOTE DE DEZVOLTARE', level=1)

doc.add_heading('10.1 Convenții de Cod', level=2)
doc.add_paragraph('Limbă: Comentariile și docstring-urile sunt în limba română', style='List Bullet')
doc.add_paragraph('Denumiri variabile: În engleză pentru compatibilitate cu biblioteci', style='List Bullet')
doc.add_paragraph('Style: PEP 8 (implicit cu Streamlit)', style='List Bullet')
doc.add_paragraph('Indentation: 4 spații', style='List Bullet')

doc.add_heading('10.2 Procesul de Debugging', level=2)
doc.add_paragraph('Pentru debugging local:')
p = doc.add_paragraph()
p.add_run('# Execuți cu verbose output\nstreamlit run exo_app.py --logger.level=debug').font.name = 'Courier New'

doc.add_heading('10.3 Adăugare Noi Dependențe', level=2)
steps = [
    'Instalează local cu pip install package_name',
    'Adaugă versiunea în requirements.txt',
    'Testează cu clean environment',
    'Commit changes la version control',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(step, style='List Number')

doc.add_heading('10.4 Gestionare Erorilor', level=2)
doc.add_paragraph(
    'Aplicația utilizează try-except blocks pentru erori în interogarile externe:'
)
p = doc.add_paragraph()
p.add_run('try:\n    data = fetch_data(input)\nexcept Exception as e:\n    st.error(f"Eroare la preluare date: {e}")').font.name = 'Courier New'

doc.add_heading('10.5 Extensii Viitoare', level=2)
extensions = [
    'Bază de date locală PostgreSQL pentru cache mai eficient',
    'Notificări email pentru noi mișiuni TESS',
    'API REST pentru integrare cu alte aplicații',
    'Machine Learning pentru clasificare candidate/false positive',
    'Suport multi-limbă pentru interfață',
    'Exportare rezultate în multiple formate (CSV, JSON, Excel)',
    'Colaborare în timp real între utilizatori',
]
for ext in extensions:
    doc.add_paragraph(ext, style='List Bullet')

doc.add_heading('10.6 Troubleshooting Comun', level=2)

doc.add_heading('Problema: Port 8501 deja în uz', level=3)
p = doc.add_paragraph()
p.add_run('streamlit run exo_app.py --server.port=8502').font.name = 'Courier New'

doc.add_heading('Problema: Import error pentru biblioteci', level=3)
p = doc.add_paragraph()
p.add_run('pip install -r requirements.txt --upgrade').font.name = 'Courier New'

doc.add_heading('Problema: Probleme cache Streamlit', level=3)
p = doc.add_paragraph()
p.add_run('streamlit cache clear').font.name = 'Courier New'

doc.add_heading('Problema: Docker build lent', level=3)
p = doc.add_paragraph()
p.add_run('# Curățare dangling images\ndocker image prune\ndocker system prune').font.name = 'Courier New'

# --- FOOTER ---
doc.add_page_break()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run('DOCUMENTAȚIE PROIECT EXOPLANETFINDER')
run.font.bold = True
run.font.size = Pt(12)

footer_para2 = doc.add_paragraph()
footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para2.add_run('Generată: Martie 2026 | Versiune: 1.0')

footer_para3 = doc.add_paragraph()
footer_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para3.add_run('Acest document conține informații tehnice complete despre arhitectura, instalare și deployment al aplicației.')

# Salvează documentul
output_path = r'c:\Exoplanete\exoplanetFinder\DOCUMENTATIE_EXOPLANETFINDER.docx'
doc.save(output_path)

print(f"✅ Documentație generată cu succes la: {output_path}")
